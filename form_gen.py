from docx import Document
import inflect
import os
import comtypes.client

# Usage: company name, date, suite #, $ amount, $ amount with commas and decimal
# All fields but the $ amount are strings

customers = [['Company, LLC', '08/30/2024', '100', 1200, '1,200.00'],
             ['Enterprise', '08/22/2024', '100', 1234.56, '1,234.56']]

# 34 entries

template_path = "1300-100 Template FINAL LIEN WAIVER.docx"

def number_to_check_words(num):
    p = inflect.engine()

    dollars, cents = f'{num:.2f}'.split(".")

    dollar_words = p.number_to_words(int(dollars)).replace(",", "")

    check_string = f'{dollar_words.capitalize()} and {cents}/100'

    return check_string

def convert_docx_to_pdf(docx_filename):
    word = comtypes.client.CreateObject("Word.Application")
    
    doc = word.Documents.Open(os.path.abspath(docx_filename))
    pdf_filename = docx_filename.replace(".docx", ".pdf")

    doc.SaveAs(os.path.abspath(pdf_filename), FileFormat=17)  # 17 = wdFormatPDF


    doc.Close()
    word.Quit()
    
    print(f"PDF saved as {pdf_filename}")

def replace_text(paragraph, placeholder, replacement):
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

doc = Document(template_path)




for customer in customers:
    doc = Document(template_path)

    for para in doc.paragraphs:
        replace_text(para, "company_name", customer[0])
        replace_text(para, "form_date", customer[1])
        replace_text(para, "suite_num", customer[2])
        replace_text(para, "payment_words", number_to_check_words(customer[3]))
        replace_text(para, "payment_num", customer[4])
        
    output_path = f'doc_files/1234-{customer[2]} {customer[0]} FINAL LIEN WAIVER.docx'
    doc.save(output_path)

    convert_docx_to_pdf(output_path)

    print(f'Generated: {output_path}')
